#!/usr/bin/env python3
"""
End-to-End Integration Example: Complete document processing pipeline
演示完整的Word/TXT文档处理管道
"""

import sys
import os
import json
import tempfile
from pathlib import Path

COMPILER_DIR = os.path.join(os.path.dirname(__file__), 'knowledge-compiler')
sys.path.insert(0, COMPILER_DIR)


def demo_complete_pipeline():
    """完整的端到端演示"""
    print("\n" + "=" * 70)
    print("END-TO-END INTEGRATION DEMO: Word/TXT Document Processing Pipeline")
    print("=" * 70)
    
    from docx import Document
    from parser.text_extractor import (
        docx_to_markdown,
        txt_to_markdown,
        split_large_text,
        has_heading_structure,
    )
    from parser.llm_structurer import _fallback_tree
    from tree_builder import parse_md_to_nodes, build_tree
    from server import detect_file_type
    
    # Create sample documents
    print("\n[STEP 1] Creating sample documents...")
    print("-" * 70)
    
    # Create Word document
    word_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc = Document()
    doc.add_heading('教育学知识体系', level=1)
    doc.add_heading('第一章 教育的基本概念', level=1)
    doc.add_heading('第一节 教育的定义', level=2)
    doc.add_paragraph('教育是人类社会一种重要现象，是指有目的、有计划、有组织地对受教育者进行思想、品德、知识和技能的培养。')
    doc.add_heading('第二节 教育的功能', level=2)
    doc.add_paragraph('教育具有多种功能：个体发展功能、社会发展功能、文化传承功能等。')
    doc.add_heading('第二章 教育与社会', level=1)
    doc.add_paragraph('教育与社会的关系是相互的。社会发展需要教育，教育发展依赖于社会。')
    word_path = word_file.name
    doc.save(word_path)
    word_file.close()
    
    print(f"✓ Created Word document: {Path(word_path).name}")
    
    # Create TXT document
    txt_file = tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8')
    txt_file.write("""第三章 学生发展

学生是教育的对象，同时也是教育过程的参与者。学生的发展是教育的根本目的。

第一节 学生身心发展的特点

学生身心发展具有多个特点：阶段性、不均衡性、互补性和个别差异性。

这些特点决定了教育必须因材施教，尊重学生的个体差异。

第二节 学生学习的规律

学生学习遵循一定的规律。认知规律、情感规律和行为规律共同指导学习过程。

教师需要理解这些规律，才能有效地组织教学活动。

第四章 教师的作用

教师是教育活动的组织者和实施者。教师的专业素养直接影响教育质量。

教师需要具备学科知识、教育理论和教学技能。同时还要具备良好的职业道德。""")
    txt_path = txt_file.name
    txt_file.close()
    
    print(f"✓ Created TXT document: {Path(txt_path).name}")
    
    # Step 2: File type detection
    print("\n[STEP 2] File type detection...")
    print("-" * 70)
    
    word_type = detect_file_type("document.docx")
    txt_type = detect_file_type("notes.txt")
    
    print(f"✓ document.docx detected as: {word_type}")
    print(f"✓ notes.txt detected as: {txt_type}")
    
    # Step 3: Text extraction
    print("\n[STEP 3] Text extraction (format conversion to Markdown)...")
    print("-" * 70)
    
    word_md = docx_to_markdown(word_path)
    print(f"✓ Word document extracted: {len(word_md)} chars")
    print(f"  Preview (first 100 chars):\n    {word_md[:100]}")
    
    txt_md = txt_to_markdown(txt_path)
    print(f"✓ TXT document extracted: {len(txt_md)} chars")
    print(f"  Preview (first 100 chars):\n    {txt_md[:100]}")
    
    # Step 4: Structure detection
    print("\n[STEP 4] Structure detection...")
    print("-" * 70)
    
    word_has_struct = has_heading_structure(word_md)
    txt_has_struct = has_heading_structure(txt_md)
    
    print(f"✓ Word document has structure: {word_has_struct}")
    print(f"  → Processing strategy: Direct markdown parsing")
    
    print(f"✓ TXT document has structure: {txt_has_struct}")
    if not txt_has_struct:
        print(f"  → Processing strategy: LLM structuring (would be applied in real scenario)")
    
    # Step 5: Text splitting
    print("\n[STEP 5] Intelligent text splitting (if needed)...")
    print("-" * 70)
    
    word_chunks = split_large_text(word_md, max_chars=1000)
    txt_chunks = split_large_text(txt_md, max_chars=1000)
    
    print(f"✓ Word document split into {len(word_chunks)} chunk(s)")
    for i, chunk in enumerate(word_chunks):
        print(f"  Chunk {i+1}: {len(chunk['text'])} chars, "
              f"first={chunk['is_first']}, last={chunk['is_last']}")
    
    print(f"✓ TXT document split into {len(txt_chunks)} chunk(s)")
    for i, chunk in enumerate(txt_chunks):
        print(f"  Chunk {i+1}: {len(chunk['text'])} chars, "
              f"first={chunk['is_first']}, last={chunk['is_last']}")
    
    # Step 6: Context bridging info
    print("\n[STEP 6] Context bridging information...")
    print("-" * 70)
    
    if len(word_chunks) > 1:
        print("Word chunks context bridging:")
        for i, chunk in enumerate(word_chunks):
            if not chunk['is_last']:
                print(f"  Chunk {i} → Chunk {i+1}: "
                      f"Previous summary will be passed as context")
    
    if len(txt_chunks) > 1:
        print("TXT chunks context bridging:")
        for i, chunk in enumerate(txt_chunks):
            if not chunk['is_last']:
                print(f"  Chunk {i} → Chunk {i+1}: "
                      f"Previous summary will be passed as context")
    
    # Step 7: Build trees from structured document
    print("\n[STEP 7] Build knowledge tree (for structured documents)...")
    print("-" * 70)
    
    # Parse first Word chunk
    if word_chunks:
        word_first_chunk = word_chunks[0]['text']
        nodes = parse_md_to_nodes(word_first_chunk)
        word_tree = build_tree(nodes)
        word_tree = {"title": "教育学知识体系", "children": word_tree}
        
        print(f"✓ Word document tree structure:")
        print(f"  Root: {word_tree['title']}")
        print(f"  Children: {len(word_tree.get('children', []))} nodes")
        
        for child in word_tree.get('children', [])[:3]:
            print(f"    - {child.get('title', 'Untitled')}")
    
    # For unstructured TXT, show fallback tree
    if txt_chunks:
        print(f"\n✓ TXT document (unstructured) - Fallback tree:")
        fallback = _fallback_tree(txt_chunks[0])
        print(f"  Root: {fallback['title']}")
        print(f"  Children: {len(fallback.get('children', []))} nodes")
        
        for child in fallback.get('children', [])[:3]:
            print(f"    - {child.get('title', 'Untitled')[:50]}")
    
    # Step 8: Show final JSON output
    print("\n[STEP 8] Final JSON output sample...")
    print("-" * 70)
    
    final_tree = {
        "title": "合并后的知识体系",
        "children": [
            word_tree if word_chunks else fallback,
            fallback if txt_chunks else word_tree,
        ]
    }
    
    print("Final tree structure (JSON):")
    print(json.dumps(final_tree, ensure_ascii=False, indent=2)[:500] + "...")
    
    # Summary
    print("\n" + "=" * 70)
    print("PROCESSING SUMMARY")
    print("=" * 70)
    
    summary = {
        "Total documents processed": 2,
        "Document 1": {
            "Type": "Word (.docx)",
            "Original size": f"{len(word_md)} chars",
            "Has structure": word_has_struct,
            "Split into chunks": len(word_chunks),
            "Processing path": "Markdown parsing" if word_has_struct else "LLM structuring",
        },
        "Document 2": {
            "Type": "Text (.txt)",
            "Original size": f"{len(txt_md)} chars",
            "Has structure": txt_has_struct,
            "Split into chunks": len(txt_chunks),
            "Processing path": "Markdown parsing" if txt_has_struct else "LLM structuring",
        },
        "Context bridging": {
            "Word context chains": len(word_chunks) - 1,
            "TXT context chains": len(txt_chunks) - 1,
            "Total context bridges": (len(word_chunks) - 1) + (len(txt_chunks) - 1),
        },
        "Output": "Unified knowledge tree JSON",
    }
    
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    
    # Cleanup
    os.unlink(word_path)
    os.unlink(txt_path)
    
    print("\n" + "=" * 70)
    print("✓ END-TO-END DEMO COMPLETED SUCCESSFULLY")
    print("=" * 70)
    print("""
Key features demonstrated:
  ✓ Automatic file type detection
  ✓ Format conversion (DOCX/TXT → Markdown)
  ✓ Structure vs unstructured text detection
  ✓ Intelligent text splitting (context-aware)
  ✓ Context bridging metadata
  ✓ Fallback tree generation
  ✓ JSON knowledge tree output

Real scenario flow:
  1. User uploads Word/TXT files
  2. System auto-detects file types
  3. Extracts text, preserving format/structure
  4. Auto-splits if too large
  5. Structured docs → direct tree parsing
  6. Unstructured docs → LLM processing (with context bridging)
  7. Final output → unified knowledge tree JSON
  8. Optional: AI enhancement (summaries, keywords, exam tips)
""")


if __name__ == "__main__":
    try:
        demo_complete_pipeline()
    except Exception as e:
        print(f"\n✗ Demo failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
