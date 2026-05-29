#!/usr/bin/env python3
"""
Advanced test: Word document processing and LLM structuring.
"""

import sys
import os
import json
import tempfile
from pathlib import Path

COMPILER_DIR = os.path.join(os.path.dirname(__file__), 'knowledge-compiler')
sys.path.insert(0, COMPILER_DIR)


def test_word_processing():
    """Test Word document extraction to markdown."""
    print("\n" + "=" * 60)
    print("TEST: Word Document Processing")
    print("=" * 60)
    
    try:
        from docx import Document
        from parser.text_extractor import docx_to_markdown
        
        # Create a test Word document
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc = Document()
            
            # Add content with different heading styles
            h1 = doc.add_heading('第一章 教育学基础', level=1)
            doc.add_paragraph('教育学是研究教育现象和教育规律的科学。')
            
            h2 = doc.add_heading('第一节 教育的定义', level=2)
            doc.add_paragraph('教育是指有目的、有计划、有组织地对受教育者进行思想、品德、知识和技能的培养。')
            
            h3 = doc.add_heading('知识点1：教育的本质', level=3)
            doc.add_paragraph('教育的本质是一种社会现象。')
            
            # Add a table
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = '概念'
            table.rows[0].cells[1].text = '定义'
            table.rows[1].cells[0].text = '教育'
            table.rows[1].cells[1].text = '有计划的培养'
            
            docx_path = tmp.name
            doc.save(docx_path)
        
        # Test extraction
        markdown = docx_to_markdown(docx_path)
        
        print(f"✓ Word document created and extracted")
        print(f"✓ Extracted {len(markdown)} characters")
        
        # Verify key content
        checks = [
            ("# 第一章 教育学基础" in markdown, "Level-1 heading"),
            ("## 第一节 教育的定义" in markdown, "Level-2 heading"),
            ("### 知识点1：教育的本质" in markdown, "Level-3 heading"),
            ("教育的本质是一种社会现象" in markdown, "Paragraph content"),
            ("|" in markdown, "Table format"),
        ]
        
        all_passed = True
        for check, desc in checks:
            status = "✓" if check else "✗"
            print(f"{status} {desc}")
            if not check:
                all_passed = False
        
        os.unlink(docx_path)
        return all_passed
        
    except Exception as e:
        print(f"✗ Word processing test failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_unstructured_text_detection():
    """Test detection of unstructured vs structured text."""
    print("\n" + "=" * 60)
    print("TEST: Structured vs Unstructured Text Detection")
    print("=" * 60)
    
    try:
        from parser.text_extractor import has_heading_structure
        
        structured = """# 主题一
内容一

## 子主题一
更多内容"""
        
        unstructured = """这是一段关于教育的讨论。教育是社会的重要组成部分。
        
教师的作用很重要。他们需要掌握学科知识和教学技能。
        
学生应该主动学习。这样才能取得好成绩。"""
        
        structured_check = has_heading_structure(structured)
        unstructured_check = has_heading_structure(unstructured)
        
        print(f"✓ Structured text detected as: {structured_check} (expected True)")
        print(f"✓ Unstructured text detected as: {unstructured_check} (expected False)")
        
        return structured_check and not unstructured_check
        
    except Exception as e:
        print(f"✗ Detection test failed: {e}")
        return False


def test_context_bridging_metadata():
    """Test that context bridging metadata is correctly set."""
    print("\n" + "=" * 60)
    print("TEST: Context Bridging Metadata Validation")
    print("=" * 60)
    
    try:
        from parser.text_extractor import split_large_text
        
        # Create a multi-section text that will be split
        text = """# 第一章 教育学基础

第一章的内容很多。 """ * 50  # Repeat to make it large
        
        chunks = split_large_text(text, max_chars=1000)
        
        print(f"✓ Text split into {len(chunks)} chunks")
        
        # Validate metadata
        issues = []
        for i, chunk in enumerate(chunks):
            # Check all required fields
            required_fields = ['text', 'index', 'total', 'is_first', 'is_last']
            for field in required_fields:
                if field not in chunk:
                    issues.append(f"Chunk {i} missing field: {field}")
            
            # Check consistency
            if chunk['index'] != i:
                issues.append(f"Chunk {i} has wrong index: {chunk['index']}")
            if chunk['total'] != len(chunks):
                issues.append(f"Chunk {i} has wrong total: {chunk['total']}")
            
            # Check first/last flags
            if i == 0 and not chunk['is_first']:
                issues.append(f"Chunk 0 should be marked first")
            if i == len(chunks) - 1 and not chunk['is_last']:
                issues.append(f"Last chunk should be marked last")
        
        if not issues:
            print("✓ All metadata fields present and correct")
            for i, chunk in enumerate(chunks):
                print(f"  Chunk {i}: {len(chunk['text'])} chars, "
                      f"first={chunk['is_first']}, last={chunk['is_last']}")
            return True
        else:
            for issue in issues:
                print(f"✗ {issue}")
            return False
            
    except Exception as e:
        print(f"✗ Metadata test failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_fallback_tree_generation():
    """Test fallback tree generation when LLM is unavailable."""
    print("\n" + "=" * 60)
    print("TEST: Fallback Tree Generation")
    print("=" * 60)
    
    try:
        from parser.llm_structurer import _fallback_tree
        
        chunk = {
            "text": """这是一个无结构的文本片段。

第一段内容很重要。我们需要理解其中的概念。

第二段介绍具体的方法。这些方法经过验证。

第三段总结要点。记住这些关键信息。""",
            "index": 0,
            "total": 1,
            "is_first": True,
            "is_last": True,
        }
        
        tree = _fallback_tree(chunk)
        
        print(f"✓ Fallback tree generated")
        print(f"✓ Root title: {tree.get('title')}")
        print(f"✓ Children count: {len(tree.get('children', []))}")
        
        # Validate structure
        if tree.get('children'):
            print("✓ Tree has children nodes")
            for i, child in enumerate(tree['children'][:3]):
                print(f"  - {child.get('title', 'Untitled')[:50]}")
            return True
        else:
            print("✗ Tree has no children")
            return False
            
    except Exception as e:
        print(f"✗ Fallback tree test failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_pipeline_with_mock_units():
    """Test the pipeline processing flow with mock units."""
    print("\n" + "=" * 60)
    print("TEST: Pipeline Flow with Mock Units")
    print("=" * 60)
    
    try:
        from parser.text_extractor import has_heading_structure, split_large_text
        
        # Simulate unit processing
        test_units = [
            {
                "name": "Structured Document",
                "text": """# 第一章 基础概念

## 第一节 定义
详细定义内容...

## 第二节 特征
特征描述内容...""",
            },
            {
                "name": "Unstructured Document",
                "text": """这是关于教育学的讨论。教育是社会的重要现象。
                
教学是教育的核心活动。教师需要掌握教学技能。
                
学生发展是教育的目标。我们需要促进学生全面发展。""",
            }
        ]
        
        for unit in test_units:
            text = unit["text"]
            has_struct = has_heading_structure(text)
            chunks = split_large_text(text)
            
            print(f"\n✓ Unit: {unit['name']}")
            print(f"  - Has structure: {has_struct}")
            print(f"  - Will use LLM: {not has_struct}")
            print(f"  - Split into {len(chunks)} chunk(s)")
            
            if not has_struct:
                print(f"  - Processing strategy: LLM structuring with context bridging")
            else:
                print(f"  - Processing strategy: Direct markdown parsing")
        
        return True
        
    except Exception as e:
        print(f"✗ Pipeline test failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_file_upload_scenario():
    """Simulate the file upload scenario."""
    print("\n" + "=" * 60)
    print("TEST: File Upload Scenario")
    print("=" * 60)
    
    try:
        from server import detect_file_type, sanitize_filename
        
        test_files = [
            "教育学原理.docx",
            "学习笔记.txt",
            "课程讲义.pdf",
        ]
        
        for filename in test_files:
            file_type = detect_file_type(filename)
            safe_name = sanitize_filename(filename)
            
            type_label = {"pdf": "PDF", "word": "Word", "text": "TXT"}.get(file_type, file_type)
            
            print(f"✓ {filename}")
            print(f"  - Detected as: {type_label}")
            print(f"  - Safe filename: {safe_name}")
            
            if file_type == "word":
                print(f"  - Will process: docx_to_markdown() → split_large_text()")
            elif file_type == "text":
                print(f"  - Will process: txt_to_markdown() → split_large_text()")
            elif file_type == "pdf":
                print(f"  - Will process: PDF split → MinerU OCR → markdown")
        
        return True
        
    except Exception as e:
        print(f"✗ Upload scenario test failed: {e}")
        return False


def main():
    """Run all advanced tests."""
    print("\n")
    print("╔" + "=" * 58 + "╗")
    print("║" + " ADVANCED DOCUMENT PROCESSING TESTS ".center(58) + "║")
    print("╚" + "=" * 58 + "╝")
    
    tests = [
        ("Word Document Processing", test_word_processing),
        ("Structured/Unstructured Detection", test_unstructured_text_detection),
        ("Context Bridging Metadata", test_context_bridging_metadata),
        ("Fallback Tree Generation", test_fallback_tree_generation),
        ("Pipeline Flow with Mock Units", test_pipeline_with_mock_units),
        ("File Upload Scenario", test_file_upload_scenario),
    ]
    
    results = {}
    for test_name, test_func in tests:
        try:
            results[test_name] = test_func()
        except Exception as e:
            print(f"\n✗ Test '{test_name}' crashed: {e}")
            import traceback
            traceback.print_exc()
            results[test_name] = False
    
    # Summary
    print("\n" + "=" * 60)
    print("ADVANCED TESTS SUMMARY")
    print("=" * 60)
    
    passed = sum(1 for v in results.values() if v)
    total = len(results)
    
    for test_name, result in results.items():
        status = "PASS ✓" if result else "FAIL ✗"
        print(f"{status}: {test_name}")
    
    print(f"\nTotal: {passed}/{total} tests passed")
    
    if passed == total:
        print("\n✓ All advanced tests passed!")
        return 0
    else:
        print(f"\n⚠ {total - passed} test(s) had issues.")
        return 1


if __name__ == "__main__":
    sys.exit(main())
